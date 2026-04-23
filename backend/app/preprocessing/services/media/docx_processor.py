"""
DOCX Processor — Media Processing Service

Purpose:
    Extract raw text from DOCX files stored on disk.
    Extracts paragraphs, tables, headers, and footers.

Input:
    local_path (str): Absolute path to the DOCX file on disk.
    mime_type  (str): MIME type of the file (expected:
                      application/vnd.openxmlformats-officedocument
                      .wordprocessingml.document).

Output:
    extracted_content (str): Raw extracted text from the DOCX.
                             Returns None on failure.

Behavior:
    - Extracts raw text only. Does NOT clean, normalize, or modify content.
    - Does NOT write to the database.
    - Does NOT modify memory_chunks.
    - Idempotent: safe to call multiple times on the same file.
    - On failure: logs error and returns None (caller handles retry).

Trigger Condition (enforced by orchestrator):
    media_files.extracted_content IS NULL
    AND mime_type = 'application/vnd.openxmlformats-officedocument
                    .wordprocessingml.document'
"""

import logging
import os
from typing import Optional

import docx

logger = logging.getLogger("echomind.preprocessing.media.docx")


def extract_text_from_docx(local_path: str, mime_type: str) -> Optional[str]:
    """
    Extract raw text from a DOCX file including paragraphs, tables,
    headers, and footers.

    Args:
        local_path: Absolute path to the DOCX file on disk.
        mime_type:  MIME type of the file.

    Returns:
        Structured extracted text as a string, or None if extraction fails.
    """
    if not os.path.exists(local_path):
        logger.error("DOCX file not found: '%s'", local_path)
        return None

    try:
        doc = docx.Document(local_path)
        parts: list[str] = []

        # Headers
        header_lines = []
        for section in doc.sections:
            for para in section.header.paragraphs:
                line = para.text.strip()
                if line:
                    header_lines.append(line)
        if header_lines:
            parts.append("[HEADER]")
            parts.extend(header_lines)

        # Paragraphs
        content_lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        if content_lines:
            parts.append("[CONTENT]")
            parts.extend(content_lines)

        # Tables
        for table in doc.tables:
            row_texts = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    row_texts.append(" | ".join(cells))
            if row_texts:
                parts.append("[TABLE]")
                parts.extend(row_texts)

        # Footers
        footer_lines = []
        for section in doc.sections:
            for para in section.footer.paragraphs:
                line = para.text.strip()
                if line:
                    footer_lines.append(line)
        if footer_lines:
            parts.append("[FOOTER]")
            parts.extend(footer_lines)

        text = "\n".join(parts).strip()
        if not text:
            return None

        logger.info("Extracted DOCX text length: %d", len(text))
        return text

    except Exception as e:
        logger.error("DOCX extraction failed for '%s': %s", local_path, e)
        return None
